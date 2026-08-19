#!/usr/bin/env python3
"""Convert the Word question banks in 115doc/ into src/data/questions.json."""

from __future__ import annotations

import json
import re
import subprocess
import tempfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "115doc"
OUTPUT_DIR = ROOT / "public" / "data"
REPORT = ROOT / "conversion-report.json"
OPTION_RE = re.compile(r"[（(]\s*([1-6])(?:\s*[）)]|(?=[\u4e00-\u9fff]))")
ANSWER_RE = re.compile(r"[1-6]")
NUMBER_RE = re.compile(r"(\d+)")
SOURCE_RE = re.compile(r"(?:\s*(?:【[^【】]*】|\[[^\[\]]*\]))+\s*$", re.S)


def clean(value: str) -> str:
    value = value.translate(str.maketrans("０１２３４５６７８９", "0123456789"))
    return re.sub(r"\s+", " ", value.replace("\u3000", " ")).strip()


def is_answer_cell(value: str) -> bool:
    return bool(re.fullmatch(r"[（(]\s*[1-6\s]+[）)]", value))


def category_from_name(path: Path) -> str:
    match = re.search(r"第\d+篇[-－](.+?)[（(]115年", path.stem)
    return match.group(1).strip() if match else path.stem


def slug(value: str) -> str:
    number = re.search(r"第(\d+)篇", value)
    return f"subject-{number.group(1)}" if number else re.sub(r"\W+", "-", value.lower()).strip("-")


def ensure_docx(source: Path, temp_dir: Path) -> Path:
    if source.suffix.lower() == ".docx":
        return source
    target = temp_dir / f"{source.stem}.docx"
    subprocess.run(
        ["textutil", "-convert", "docx", "-output", str(target), str(source)],
        check=True,
        capture_output=True,
        text=True,
    )
    return target


def answer_digits(value: str) -> list[str]:
    return list(dict.fromkeys(ANSWER_RE.findall(value)))


def parse_question_text(raw: str) -> tuple[str, list[dict[str, str]], str] | None:
    raw = clean(raw)
    if not raw:
        return None

    source_match = SOURCE_RE.search(raw)
    explanation = clean(source_match.group()) if source_match else ""
    body = raw[: source_match.start()].strip() if source_match else raw
    matches = list(OPTION_RE.finditer(body))
    if len(matches) < 2:
        return None

    # The first true option marker should start a sequential 1, 2, 3... run.
    start = next((i for i, match in enumerate(matches) if match.group(1) == "1"), None)
    if start is None:
        return None
    run = [matches[start]]
    expected = 2
    for match in matches[start + 1 :]:
        digit = int(match.group(1))
        if digit == expected:
            run.append(match)
            expected += 1
            if len(run) == 3:
                break
        elif digit > expected and len(run) >= 2:
            break
    if len(run) < 2:
        return None

    question = clean(body[: run[0].start()]).rstrip("：:。")
    options: list[dict[str, str]] = []
    for index, marker in enumerate(run):
        end = run[index + 1].start() if index + 1 < len(run) else len(body)
        text = clean(body[marker.end() : end]).rstrip("。")
        if text:
            options.append({"id": chr(65 + int(marker.group(1)) - 1), "text": text})
    if not question or len(options) < 2:
        return None
    return question, options, explanation


def unique_cells(cells: list[str]) -> list[str]:
    result: list[str] = []
    for value in cells:
        value = clean(value)
        if value and value not in result:
            result.append(value)
    return result


def records_from_table(document: Document) -> list[tuple[str, str, str]]:
    records: list[tuple[str, str, str]] = []
    for table in document.tables:
        for row in table.rows:
            cells = unique_cells([cell.text for cell in row.cells])
            if len(cells) < 3:
                continue
            answer = cells[0]
            number = next((value for value in cells[1:] if re.fullmatch(r"\s*\d+\.?\s*", value)), "")
            candidates = [value for value in cells[1:] if value != number]
            question = max(candidates, key=len, default="")
            if answer_digits(answer) and number and OPTION_RE.search(question):
                records.append((answer, number, question))
    return records


def records_from_paragraphs(document: Document) -> list[tuple[str, str, str]]:
    values: list[str] = []
    for paragraph in document.paragraphs:
        value = clean(paragraph.text)
        if not value:
            continue
        merged = re.match(r"^([（(]\s*[1-6\s]+[）)])\s*(\d+\.)\s*(.+)$", value)
        values.extend(merged.groups() if merged else [value])
    records: list[tuple[str, str, str]] = []
    index = 0
    while index + 2 < len(values):
        answer, number = values[index : index + 2]
        if is_answer_cell(answer) and re.fullmatch(r"\d+\.?", number):
            end = index + 2
            while end < len(values):
                if end + 1 < len(values) and is_answer_cell(values[end]) and re.fullmatch(r"\d+\.?", values[end + 1]):
                    break
                end += 1
            question = " ".join(values[index + 2 : end])
            if OPTION_RE.search(question):
                records.append((answer, number, question))
            index = end
            continue
        index += 1
    return records


def records_from_unnumbered_paragraphs(document: Document) -> list[tuple[str, str, str]]:
    """Some legacy files contain answer/question pairs without visible question numbers."""
    values = [clean(paragraph.text) for paragraph in document.paragraphs if clean(paragraph.text)]
    records: list[tuple[str, str, str]] = []
    index = 0
    question_number = 1
    while index < len(values):
        answer = values[index]
        if is_answer_cell(answer) and index + 1 < len(values):
            question = values[index + 1]
            index += 2
            if index < len(values) and re.fullmatch(r"[【\[].*[】\]]", values[index]):
                question = f"{question} {values[index]}"
                index += 1
            if OPTION_RE.search(question):
                records.append((answer, f"{question_number}.", question))
                question_number += 1
                continue
        index += 1
    return records


def parse_file(source: Path, temp_dir: Path) -> tuple[list[dict], list[dict]]:
    docx_path = ensure_docx(source, temp_dir)
    document = Document(docx_path)
    category = category_from_name(source)
    category_id = slug(source.stem)
    records = records_from_table(document) or records_from_paragraphs(document) or records_from_unnumbered_paragraphs(document)
    questions: list[dict] = []
    errors: list[dict] = []

    for answer_raw, number_raw, question_raw in records:
        number_match = NUMBER_RE.search(number_raw)
        parsed = parse_question_text(question_raw)
        answers = [chr(64 + int(value)) for value in answer_digits(answer_raw)]
        if not number_match or not parsed:
            errors.append({"source": source.name, "number": number_raw, "reason": "無法解析題目或選項"})
            continue
        question, options, explanation = parsed
        option_ids = {option["id"] for option in options}
        if not answers or not set(answers).issubset(option_ids):
            errors.append({"source": source.name, "number": number_raw, "reason": f"答案 {answers} 不在選項 {sorted(option_ids)} 中"})
            continue
        number = int(number_match.group(1))
        questions.append({
            "id": f"{category_id}-{number:03d}",
            "category": category,
            "number": number,
            "type": "multiple" if len(answers) > 1 else "single",
            "question": question,
            "options": options,
            "answers": answers,
            "explanation": explanation,
            "source": source.name,
        })
    return questions, errors


def main() -> None:
    sources = sorted(
        [path for path in SOURCE_DIR.iterdir() if path.suffix.lower() in {".doc", ".docx"} and not path.name.startswith(".~")],
        key=lambda path: int(re.search(r"第(\d+)篇", path.name).group(1)),
    )
    all_questions: list[dict] = []
    all_errors: list[dict] = []
    file_stats: list[dict] = []

    with tempfile.TemporaryDirectory(prefix="exam115-") as temp:
        temp_dir = Path(temp)
        for source in sources:
            questions, errors = parse_file(source, temp_dir)
            all_questions.extend(questions)
            all_errors.extend(errors)
            file_stats.append({"source": source.name, "parsed": len(questions), "errors": len(errors)})
            print(f"{source.name}: {len(questions)} 題，{len(errors)} 個錯誤")

    counts = Counter(question["category"] for question in all_questions)
    categories = []
    for source in sources:
        name = category_from_name(source)
        subject_questions = [q for q in all_questions if q["category"] == name]
        categories.append({
            "id": slug(source.stem),
            "name": name,
            "count": counts[name],
            "singleCount": sum(q["type"] == "single" for q in subject_questions),
            "multipleCount": sum(q["type"] == "multiple" for q in subject_questions),
        })
    duplicate_ids = [item for item, count in Counter(q["id"] for q in all_questions).items() if count > 1]
    payload = {
        "generatedAt": datetime.now(timezone.utc).isoformat(),
        "questionCount": len(all_questions),
        "categories": categories,
        "questions": all_questions,
    }
    report = {
        "generatedAt": payload["generatedAt"],
        "sourceFiles": len(sources),
        "questionCount": len(all_questions),
        "duplicateIds": duplicate_ids,
        "files": file_stats,
        "errors": all_errors,
    }
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    index_payload = {key: payload[key] for key in ("generatedAt", "questionCount", "categories")}
    (OUTPUT_DIR / "index.json").write_text(json.dumps(index_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    for category in categories:
        subject_questions = [q for q in all_questions if q["category"] == category["name"]]
        subject_payload = {"category": category, "questionCount": len(subject_questions), "questions": subject_questions}
        (OUTPUT_DIR / f"{category['id']}.json").write_text(json.dumps(subject_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"\n已輸出 {OUTPUT_DIR.relative_to(ROOT)}：1 份索引、{len(categories)} 份科目題庫，共 {len(all_questions)} 題")
    print(f"驗證報告：{REPORT.relative_to(ROOT)}")
    if duplicate_ids or all_errors:
        print(f"注意：{len(duplicate_ids)} 個重複 ID，{len(all_errors)} 個解析錯誤")


if __name__ == "__main__":
    main()
